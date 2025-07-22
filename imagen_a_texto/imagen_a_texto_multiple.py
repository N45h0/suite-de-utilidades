# Script para convertir varias imágenes a texto y guardar en DOCX
# Uso: Ejecutar y seguir la interfaz gráfica

import pytesseract
from PIL import Image
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import os

pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

class OCRBatchApp:
    def __init__(self, root):
        self.root = root
        self.root.title("OCR Batch - Español a DOCX")
        self.image_paths = []
        self.combined_mode = tk.BooleanVar(value=True)
        self.create_widgets()
    def create_widgets(self):
        tk.Label(self.root, text="Selecciona imágenes y elige el modo:").pack(pady=5)
        self.btn_select = tk.Button(self.root, text="Seleccionar imágenes", command=self.select_images)
        self.btn_select.pack(pady=5)
        tk.Checkbutton(self.root, text="Combinar todo en un solo DOCX", variable=self.combined_mode).pack(pady=5)
        self.btn_convert = tk.Button(self.root, text="Convertir imágenes", command=self.convert_images, state=tk.DISABLED)
        self.btn_convert.pack(pady=5)
        self.status_label = tk.Label(self.root, text="", wraplength=400)
        self.status_label.pack(pady=10)
    def select_images(self):
        self.image_paths = filedialog.askopenfilenames(
            filetypes=[("Imágenes", "*.png;*.jpg;*.jpeg;*.bmp;*.tiff")]
        )
        if self.image_paths:
            self.btn_convert.config(state=tk.NORMAL)
    def convert_images(self):
        if self.combined_mode.get():
            doc = Document()
            for path in self.image_paths:
                img = Image.open(path)
                text = pytesseract.image_to_string(img, lang="spa")
                doc.add_paragraph(f"Archivo: {os.path.basename(path)}")
                doc.add_paragraph(text)
                doc.add_page_break()
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
            if save_path:
                doc.save(save_path)
                self.status_label.config(text="Guardado en: " + save_path)
        else:
            for path in self.image_paths:
                img = Image.open(path)
                text = pytesseract.image_to_string(img, lang="spa")
                doc = Document()
                doc.add_paragraph(text)
                save_path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=os.path.splitext(os.path.basename(path))[0]+".docx", filetypes=[("Word Document", "*.docx")])
                if save_path:
                    doc.save(save_path)
            self.status_label.config(text="Archivos guardados.")

if __name__ == "__main__":
    root = tk.Tk()
    app = OCRBatchApp(root)
    root.mainloop()
