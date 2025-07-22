import pytesseract
from PIL import Image
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import os

# Configurar Tesseract si es necesario
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

class OCRBatchApp:
    def __init__(self, root):
        self.root = root
        self.root.title("OCR Batch - Español a DOCX")
        self.image_paths = []
        self.combined_mode = tk.BooleanVar(value=True)  # Modo combinado por defecto
        
        self.create_widgets()
    
    def create_widgets(self):
        # Botones y opciones
        tk.Label(self.root, text="Selecciona imágenes y elige el modo:").pack(pady=5)
        
        self.btn_select = tk.Button(self.root, text="Seleccionar imágenes", command=self.select_images)
        self.btn_select.pack(pady=5)
        
        tk.Checkbutton(self.root, text="Combinar todo en un solo DOCX", variable=self.combined_mode).pack(pady=5)
        
        self.btn_convert = tk.Button(self.root, text="Convertir imágenes", command=self.convert_images, state=tk.DISABLED)
        self.btn_convert.pack(pady=5)
        
        self.status_label = tk.Label(self.root, text="", wraplength=400)
        self.status_label.pack(pady=10)
    
    def select_images(self):
        self.image_paths = filedialog.askopenfilenames(
            filetypes=[("Imágenes", "*.png;*.jpg;*.jpeg;*.bmp;*.tiff")]
        )
        if self.image_paths:
            self.btn_convert.config(state=tk.NORMAL)
            self.status_label.config(text=f"{len(self.image_paths)} imágenes seleccionadas")
    
    def convert_images(self):
        try:
            doc = Document()
            output_text = ""
            
            for i, image_path in enumerate(self.image_paths, 1):
                img = Image.open(image_path)
                text = pytesseract.image_to_string(img, lang='spa')
                
                if self.combined_mode.get():
                    doc.add_paragraph(f"--- Imagen {i} ---\n{text}\n")
                else:
                    individual_doc = Document()
                    individual_doc.add_paragraph(text)
                    output_dir = filedialog.askdirectory(title="Guardar documentos individuales en...")
                    if output_dir:
                        output_path = os.path.join(output_dir, f"documento_{i}.docx")
                        individual_doc.save(output_path)
            
            if self.combined_mode.get():
                output_file = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Documentos Word", "*.docx")]
                )
                if output_file:
                    doc.save(output_file)
                    self.status_label.config(text=f"Archivo combinado guardado:\n{output_file}")
                    messagebox.showinfo("Éxito", f"{len(self.image_paths)} imágenes procesadas")
            else:
                self.status_label.config(text=f"{len(self.image_paths)} documentos guardados")
            
        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = OCRBatchApp(root)
    root.mainloop()