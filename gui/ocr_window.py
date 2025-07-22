import customtkinter as ctk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import pytesseract
from docx import Document
import os
import threading

# Configurar la ruta de Tesseract para Windows
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

class OcrWindow(ctk.CTkToplevel):
    def __init__(self, master=None):
        super().__init__(master)
        self.title("Extractor de Texto de Imágenes (OCR)")
        self.geometry("800x600")
        self.configure(fg_color="#f8f9fa")
        self.transient(master)
        self.grab_set()

        # --- Layout Principal ---
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # --- Contenedor Principal ---
        main_container = ctk.CTkFrame(self, fg_color="white", corner_radius=10, border_width=1, border_color="#e5e7eb")
        main_container.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")
        main_container.grid_columnconfigure(0, weight=1)
        main_container.grid_rowconfigure(1, weight=1)

        # --- Cabecera ---
        header_frame = ctk.CTkFrame(main_container, fg_color="transparent")
        header_frame.grid(row=0, column=0, padx=20, pady=20, sticky="ew")
        
        title_label = ctk.CTkLabel(header_frame, text="Extractor de Texto de Imágenes (OCR)", font=ctk.CTkFont(size=24, weight="bold"))
        title_label.pack(anchor="w")
        
        subtitle_label = ctk.CTkLabel(header_frame, text="Extrae texto de imágenes individuales o procesa una carpeta completa.", text_color="#22292f")
        subtitle_label.pack(anchor="w")

        # --- Pestañas ---
        self.tab_view = ctk.CTkTabview(main_container, fg_color="transparent")
        self.tab_view.grid(row=1, column=0, padx=20, pady=10, sticky="nsew")

        self.tab_individual = self.tab_view.add("Imagen Individual")
        self.tab_folder = self.tab_view.add("Procesar Carpeta")

        self.setup_individual_tab()
        self.setup_folder_tab()

    def setup_individual_tab(self):
        self.tab_individual.grid_columnconfigure((0, 1), weight=1)
        self.tab_individual.grid_rowconfigure(0, weight=1)

        # --- Columna Izquierda (Imagen) ---
        left_frame = ctk.CTkFrame(self.tab_individual, fg_color="transparent")
        left_frame.grid(row=0, column=0, padx=(0, 10), sticky="nsew")
        left_frame.grid_rowconfigure(1, weight=1)

        select_file_button = ctk.CTkButton(left_frame, text="Seleccionar Archivo", command=self.select_image_file)
        select_file_button.grid(row=0, column=0, sticky="ew", pady=(0, 10))

        self.image_preview_frame = ctk.CTkFrame(left_frame, fg_color="#f3f4f6", border_width=2, border_color="#e5e7eb", corner_radius=5)
        self.image_preview_frame.grid(row=1, column=0, sticky="nsew")
        self.image_preview_label = ctk.CTkLabel(self.image_preview_frame, text="Vista previa de la imagen")
        self.image_preview_label.pack(expand=True)

        # --- Columna Derecha (Texto) ---
        right_frame = ctk.CTkFrame(self.tab_individual, fg_color="transparent")
        right_frame.grid(row=0, column=1, padx=(10, 0), sticky="nsew")
        right_frame.grid_rowconfigure(1, weight=1)

        text_label = ctk.CTkLabel(right_frame, text="Texto Extraído", font=ctk.CTkFont(weight="bold"))
        text_label.grid(row=0, column=0, sticky="w")

        self.text_output = ctk.CTkTextbox(right_frame, state="disabled", fg_color="#f3f4f6", text_color="#111418")
        self.text_output.grid(row=1, column=0, sticky="nsew", pady=(5, 10))

        copy_button = ctk.CTkButton(right_frame, text="Copiar Texto", command=self.copy_text)
        copy_button.grid(row=2, column=0, sticky="ew")

        save_docx_button = ctk.CTkButton(right_frame, text="Guardar como DOCX", command=self.save_docx)
        save_docx_button.grid(row=3, column=0, sticky="ew", pady=(5,0))

        self.selected_image_path = None
        self.last_ocr_text = ""

    def select_image_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg *.jpeg *.png *.bmp *.tiff")])
        if not file_path:
            return
        self.selected_image_path = file_path
        self.display_image(file_path)
        self.text_output.configure(state="normal")
        self.text_output.delete("1.0", "end")
        self.text_output.insert("1.0", "Procesando OCR...")
        self.text_output.configure(state="disabled")
        threading.Thread(target=self.run_ocr_image, args=(file_path,)).start()

    def run_ocr_image(self, file_path):
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="spa")
            self.last_ocr_text = text
            self.after(0, lambda: self.show_ocr_result(text))
        except Exception as e:
            self.after(0, lambda: self.show_ocr_result(f"Error: {e}"))

    def show_ocr_result(self, text):
        self.text_output.configure(state="normal", text_color="#111418")
        self.text_output.delete("1.0", "end")
        self.text_output.insert("1.0", text)
        self.text_output.configure(state="disabled")

    def display_image(self, file_path):
        try:
            img = Image.open(file_path)
            img.thumbnail((self.image_preview_frame.winfo_width() - 20, self.image_preview_frame.winfo_height() - 20))
            photo = ImageTk.PhotoImage(img)
            self.image_preview_label.configure(image=photo, text="")
            self.image_preview_label.image = photo
        except Exception as e:
            self.image_preview_label.configure(text=f"Error al cargar imagen:\n{e}")

    def copy_text(self):
        text = self.text_output.get("1.0", "end-1c")
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)
            messagebox.showinfo("Copiado", "Texto copiado al portapapeles.")

    def save_docx(self):
        if not self.last_ocr_text.strip():
            messagebox.showwarning("Sin texto", "No hay texto extraído para guardar.")
            return
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if save_path:
            doc = Document()
            doc.add_paragraph(self.last_ocr_text)
            doc.save(save_path)
            messagebox.showinfo("Guardado", f"Archivo guardado en: {save_path}")

    def setup_folder_tab(self):
        self.tab_folder.grid_columnconfigure(0, weight=1)

        folder_label = ctk.CTkLabel(self.tab_folder, text="Selecciona una carpeta para procesar todas sus imágenes.")
        folder_label.grid(row=0, column=0, padx=10, pady=10, sticky="w")

        select_folder_button = ctk.CTkButton(self.tab_folder, text="Seleccionar Carpeta", command=self.select_folder)
        select_folder_button.grid(row=1, column=0, padx=10, pady=10, sticky="w")

        self.selected_folder_label = ctk.CTkLabel(self.tab_folder, text="Carpeta no seleccionada", text_color="gray")
        self.selected_folder_label.grid(row=2, column=0, padx=10, pady=5, sticky="w")

        process_button = ctk.CTkButton(self.tab_folder, text="Iniciar Proceso", command=self.process_folder)
        process_button.grid(row=3, column=0, padx=10, pady=20, sticky="w")

        self.progress_bar_folder = ctk.CTkProgressBar(self.tab_folder)
        self.progress_bar_folder.set(0)
        self.progress_bar_folder.grid(row=4, column=0, padx=10, pady=5, sticky="ew")

        self.status_label_folder = ctk.CTkLabel(self.tab_folder, text="Listo.")
        self.status_label_folder.grid(row=5, column=0, padx=10, pady=5, sticky="w")

        self.folder_processing = False

    def select_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.selected_folder_label.configure(text=folder_path)

    def process_folder(self):
        folder_path = self.selected_folder_label.cget("text")
        if not os.path.isdir(folder_path):
            self.status_label_folder.configure(text="Por favor, selecciona una carpeta válida.", text_color="orange")
            return
        if self.folder_processing:
            return
        self.folder_processing = True
        self.status_label_folder.configure(text="Procesando carpeta...", text_color="black")
        self.progress_bar_folder.set(0)
        threading.Thread(target=self.run_ocr_folder, args=(folder_path,)).start()

    def run_ocr_folder(self, folder_path):
        try:
            image_files = [f for f in os.listdir(folder_path) if f.lower().endswith((".jpg", ".jpeg", ".png", ".bmp", ".tiff"))]
            if not image_files:
                self.after(0, lambda: self.status_label_folder.configure(text="No se encontraron imágenes en la carpeta.", text_color="orange"))
                self.folder_processing = False
                return
            doc = Document()
            for i, filename in enumerate(image_files, 1):
                file_path = os.path.join(folder_path, filename)
                try:
                    img = Image.open(file_path)
                    text = pytesseract.image_to_string(img, lang="spa")
                    doc.add_paragraph(f"Archivo: {filename}")
                    doc.add_paragraph(text)
                    doc.add_page_break()
                except Exception as e:
                    doc.add_paragraph(f"Error procesando {filename}: {e}")
                progress = i / len(image_files)
                self.after(0, lambda p=progress: self.progress_bar_folder.set(p))
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
            if save_path:
                doc.save(save_path)
                self.after(0, lambda: self.status_label_folder.configure(text=f"Guardado en: {save_path}", text_color="green"))
            else:
                self.after(0, lambda: self.status_label_folder.configure(text="Proceso cancelado.", text_color="orange"))
        except Exception as e:
            self.after(0, lambda: self.status_label_folder.configure(text=f"Error: {e}", text_color="red"))
        self.folder_processing = False


if __name__ == '__main__':
    app = ctk.CTk()
    app.withdraw()
    window = OcrWindow(app)
    app.mainloop()
